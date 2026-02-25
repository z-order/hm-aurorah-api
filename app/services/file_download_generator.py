"""
File download generator service

Stateless file generation service for converting segment data into
downloadable document formats (TXT, DOCX, PDF, XLSX).

No database access -- pure data-in, bytes-out.

Segment structure:
    [{"sid": 1, "text": "Hello"}, {"sid": 2, "text": "World"}]

MIME Types:
    TXT:  text/plain; charset=utf-8
    DOCX: application/vnd.openxmlformats-officedocument.wordprocessingml.document
    PDF:  application/pdf
    XLSX: application/vnd.openxmlformats-officedocument.spreadsheetml.sheet
"""

import io
import logging
from typing import Any, cast

from app.core.logger import get_logger

logger = get_logger(__name__, logging.INFO)


# MIME type constants for each supported file format
MIME_TYPES: dict[str, str] = {
    "txt": "text/plain; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

# File extension constants for Content-Disposition header
FILE_EXTENSIONS: dict[str, str] = {
    "txt": ".txt",
    "docx": ".docx",
    "pdf": ".pdf",
    "xlsx": ".xlsx",
}


def generate_txt(segments: list[dict[str, Any]], include_numbers: bool) -> bytes:
    """
    Generate a plain text file from segments.

    Format with numbers:    "[1] Hello\\n[2] World\\n"
    Format without numbers: "Hello\\nWorld\\n"

    Args:
        segments: List of segment dicts with "sid" and "text" keys
        include_numbers: Whether to prefix each line with [sid]

    Returns:
        UTF-8 encoded bytes of the text file
    """
    lines: list[str] = []
    for seg in segments:
        if include_numbers:
            lines.append(f"[{seg['sid']}] {seg['text']}")
        else:
            lines.append(seg["text"])

    content = "\n".join(lines) + "\n"
    return content.encode("utf-8")


def generate_docx(segments: list[dict[str, Any]], include_numbers: bool, column_title: str) -> bytes:
    """
    Generate a DOCX file from segments using python-docx.

    Each segment becomes a paragraph, optionally prefixed with [sid].
    The document title is set from column_title.

    Args:
        segments: List of segment dicts with "sid" and "text" keys
        include_numbers: Whether to prefix each paragraph with [sid]
        column_title: Document title displayed as the heading

    Returns:
        DOCX file bytes
    """
    from docx import Document  # type: ignore[import-untyped]

    document = cast(Any, Document())
    document.add_heading(column_title, level=1)

    for seg in segments:
        if include_numbers:
            text = f"[{seg['sid']}] {seg['text']}"
        else:
            text = seg["text"]
        document.add_paragraph(text)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


# Noto font directories (fonts-noto-core + fonts-noto-cjk)
_NOTO_FONT_DIRS: list[str] = [
    "/usr/share/fonts/truetype/noto",
    "/usr/share/fonts/opentype/noto",
]

# Primary font: NotoSans covers Latin, Cyrillic, Greek, and many more.
_PRIMARY_FONT_NAME = "NotoSans-Regular.ttf"

# CJK font (separate .ttc collection)
_CJK_FONT_NAME = "NotoSansCJK-Regular.ttc"


def _discover_noto_fonts() -> tuple[str | None, str | None, list[str]]:
    """
    Discover installed Noto Sans fonts for PDF generation.

    Returns:
        (primary_path, cjk_path, extra_paths) where extra_paths are
        all other NotoSans*-Regular.ttf files found (for fallback).
    """
    from pathlib import Path

    primary: str | None = None
    cjk: str | None = None
    extras: list[str] = []

    for font_dir in _NOTO_FONT_DIRS:
        d = Path(font_dir)
        if not d.is_dir():
            continue

        for f in sorted(d.iterdir()):
            name = f.name
            if name == _PRIMARY_FONT_NAME:
                primary = str(f)
            elif name == _CJK_FONT_NAME:
                cjk = str(f)
            elif name.startswith("NotoSans") and name.endswith("-Regular.ttf"):
                extras.append(str(f))

    return primary, cjk, extras


def generate_pdf(segments: list[dict[str, Any]], include_numbers: bool, column_title: str) -> bytes:
    """
    Generate a PDF file from segments using fpdf2.

    Each segment becomes a multi-line cell, optionally prefixed with [sid].
    Uses Noto Sans as primary font with CJK + per-script fallbacks for
    full multilingual support. Requires system packages:
        apt-get install -y fonts-noto-core fonts-noto-cjk

    Args:
        segments: List of segment dicts with "sid" and "text" keys
        include_numbers: Whether to prefix each line with [sid]
        column_title: Document title displayed at the top

    Returns:
        PDF file bytes
    """
    from fpdf import FPDF  # type: ignore[import-untyped]

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    primary_path, cjk_path, extra_paths = _discover_noto_fonts()

    fallback_families: list[str] = []

    if primary_path:
        pdf.add_font("NotoSans", "", primary_path)
    if cjk_path:
        pdf.add_font("NotoSansCJK", "", cjk_path)
        fallback_families.append("NotoSansCJK")

    for fpath in extra_paths:
        from pathlib import Path
        family = Path(fpath).stem  # e.g. "NotoSansArabic-Regular" -> stem
        pdf.add_font(family, "", fpath)
        fallback_families.append(family)

    if fallback_families:
        pdf.set_fallback_fonts(fallback_families, exact_match=False)

    try:
        pdf.set_text_shaping(True)
    except Exception:
        pass

    main_font = "NotoSans" if primary_path else "Helvetica"

    # Title
    pdf.set_font(main_font, size=16)
    pdf.cell(0, 10, column_title, new_x="LMARGIN", new_y="NEXT")

    # Body
    pdf.set_font(main_font, size=10)
    pdf.ln(5)

    for seg in segments:
        if include_numbers:
            text = f"[{seg['sid']}] {seg['text']}"
        else:
            text = seg["text"]
        pdf.multi_cell(0, 6, text)  # type: ignore[no-untyped-call]
        pdf.ln(2)

    return bytes(pdf.output())


def generate_xlsx(segments: list[dict[str, Any]], include_numbers: bool, column_title: str) -> bytes:
    """
    Generate an XLSX file from segments using openpyxl.

    When include_numbers is True:
        Column A = segment number (sid), Column B = translated text
    When include_numbers is False:
        Column A = translated text only

    The first row is a header row. Sheet title is set from column_title.

    Args:
        segments: List of segment dicts with "sid" and "text" keys
        include_numbers: Whether to include a segment number column
        column_title: Worksheet title (also used for header)

    Returns:
        XLSX file bytes
    """
    from openpyxl import Workbook  # type: ignore[import-untyped]

    wb = Workbook()
    ws = wb.active
    ws.title = column_title  # type: ignore[union-attr]

    # Write header row
    if include_numbers:
        ws.append(["#", column_title])  # type: ignore[union-attr]
    else:
        ws.append([column_title])  # type: ignore[union-attr]

    # Write segment rows
    for seg in segments:
        if include_numbers:
            ws.append([seg["sid"], seg["text"]])  # type: ignore[union-attr]
        else:
            ws.append([seg["text"]])  # type: ignore[union-attr]

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.read()
